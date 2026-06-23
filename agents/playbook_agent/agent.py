from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x1A, 0x4F, 0x72)
GOLD  = RGBColor(0xC9, 0x92, 0x2A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x1A, 0x2C, 0x3D)
GREY  = RGBColor(0x4A, 0x65, 0x80)
GREEN_DARK  = RGBColor(0x0D, 0x6E, 0x2C)
GREEN_LIGHT = RGBColor(0x16, 0x63, 0x34)
STEEL = RGBColor(0x2E, 0x6A, 0x96)
RED   = RGBColor(0xDC, 0x26, 0x26)

doc = Document()

section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)


def shade_paragraph(para, hex_fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)


def shade_cell(cell, hex_fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def add_heading(text, size=18, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_subheading(text, color=GOLD, size=11, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_body(text, size=10.5, color=DARK, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_bullet(text, size=10.5, color=DARK):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C9922A')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_screenshot_placeholder(label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(8)
    shade_paragraph(p, 'E8F0F7')
    run = p.add_run('  ' + label + '  ')
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = STEEL


# ═══════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════
cover_bg = doc.add_paragraph()
shade_paragraph(cover_bg, '1A4F72')
cover_bg.paragraph_format.space_before = Pt(0)
cover_bg.paragraph_format.space_after  = Pt(0)

p = doc.add_paragraph()
shade_paragraph(p, '1A4F72')
run = p.add_run('CHRYSELYS')
run.bold = True; run.font.size = Pt(30); run.font.color.rgb = GOLD
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(56)
p.paragraph_format.space_after  = Pt(4)

p2 = doc.add_paragraph()
shade_paragraph(p2, '1A4F72')
run2 = p2.add_run('Forecasting Accelerator')
run2.bold = True; run2.font.size = Pt(22); run2.font.color.rgb = WHITE
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(8)

p3 = doc.add_paragraph()
shade_paragraph(p3, '1A4F72')
run3 = p3.add_run('Sales Playbook')
run3.italic = True; run3.font.size = Pt(15); run3.font.color.rgb = RGBColor(0xA8, 0xC4, 0xD4)
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(4)

p4 = doc.add_paragraph()
shade_paragraph(p4, '1A4F72')
run4 = p4.add_run('For Sales & Business Development Use Only')
run4.font.size = Pt(10); run4.font.color.rgb = RGBColor(0xA8, 0xC4, 0xD4)
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_after = Pt(56)

doc.add_page_break()

# ═══════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════
add_heading('Table of Contents', size=16)
add_divider()
toc_entries = [
    ('1', 'Product Overview', '3'),
    ('2', 'Target Audience & Ideal Customer Profile', '4'),
    ('3', 'Feature Deep Dive', '5'),
    ('4', 'Recommended Demo Flow', '6'),
    ('5', 'Objection Handling Guide', '7'),
    ('6', 'Competitive Positioning', '8'),
    ('7', 'Sample Forecast Output', '9'),
    ('8', 'Discovery Questions', '10'),
    ('9', 'Call to Action & Closing Moves', '11'),
]
for num, title, pg in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run_n = p.add_run(f'{num}.  ')
    run_n.bold = True; run_n.font.color.rgb = GOLD; run_n.font.size = Pt(11)
    run_t = p.add_run(title)
    run_t.font.color.rgb = DARK; run_t.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 1. PRODUCT OVERVIEW
# ═══════════════════════════════════════════════
add_heading('1. Product Overview', size=17)
add_divider()
add_body(
    'Chryselys Forecasting Accelerator is an AI-powered commercial forecasting platform purpose-built '
    'for pharmaceutical and biotech teams. It combines a rigorous epidemiology-driven patient funnel model '
    'with a conversational AI copilot, enabling commercial, market access, and strategy teams to build '
    'credible, data-backed revenue forecasts in a fraction of the time of legacy spreadsheet workflows.'
)

add_subheading('Elevator Pitch (30 seconds)')
add_body(
    '"Forecasting Accelerator turns weeks of spreadsheet work into minutes. You enter a product, an indication, '
    'and a geography and our AI builds a full patient funnel forecast with benchmarked assumptions and '
    'evidence-backed rationale, ready for your brand plan or investor deck. No more manual literature hunts. '
    'No more broken Excel models. Just clear, defensible forecasts."'
)

add_subheading('Core Value Proposition')
for b in [
    'Speed - From blank page to publish-ready forecast in under 30 minutes.',
    'Credibility - Every assumption backed by published literature or real-world data sources.',
    'Collaboration - Shared AI copilot that any team member can query in plain language.',
    'Flexibility - Supports 9 countries, 6 major indication templates, and fully custom funnels.',
    'Auditability - Full assumption log with sources for regulatory or investor review.',
]:
    add_bullet(b)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 2. TARGET AUDIENCE & ICP
# ═══════════════════════════════════════════════
add_heading('2. Target Audience & Ideal Customer Profile', size=17)
add_divider()

add_subheading('Primary Buyers')
icp_rows = [
    ('Persona', 'Title Examples', 'Key Pain Point'),
    ('Commercial Strategy', 'VP Commercial, Head of Strategy', 'Forecast cycle takes 6-8 weeks; assumptions undocumented'),
    ('Market Access', 'Director Market Access, HEOR Lead', 'Budget impact models are slow to build and update'),
    ('Finance / FP&A', 'CFO, Commercial Finance Director', 'Revenue models lack epidemiology grounding; hard to defend'),
    ('Medical Affairs', 'Head of Medical Affairs, MSL Lead', 'No structured link between clinical evidence and market estimates'),
    ('Consulting / Agencies', 'Engagement Manager, Principal', 'Client deliverables require defensible, transparent assumptions'),
]
tbl = doc.add_table(rows=len(icp_rows), cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
c_widths = [Inches(1.8), Inches(2.0), Inches(2.6)]
for i, row_data in enumerate(icp_rows):
    row = tbl.rows[i]
    for j, (cell, w) in enumerate(zip(row.cells, c_widths)):
        cell.width = w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(row_data[j])
        run.font.size = Pt(9.5)
        if i == 0:
            run.bold = True; run.font.color.rgb = WHITE
            shade_cell(cell, '1A4F72')
        else:
            run.font.color.rgb = DARK
            shade_cell(cell, 'F5F6F8' if i % 2 == 0 else 'FFFFFF')
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

add_body('')
add_subheading('Target Company Profile')
for b in [
    'Pharma / biotech companies with products in late-stage development, launch, or post-launch phases',
    'Companies operating in oncology, immunology, neurology, cardiology, or rare disease',
    'Teams spending >2 weeks per forecast cycle on epidemiology research and model building',
    'Organizations preparing investor materials, brand plans, or payer submissions',
    'Management consulting firms and market research agencies serving pharma clients',
]:
    add_bullet(b)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 3. FEATURE DEEP DIVE
# ═══════════════════════════════════════════════
add_heading('3. Feature Deep Dive', size=17)
add_divider()

features = [
    ('AI Copilot Chat',
     'An embedded AI assistant that answers questions in plain English: parameter explanations, benchmark ranges, '
     'scenario interpretation, and methodology questions. It cites published literature and provides rationale behind every number.'),
    ('Patient Funnel Engine',
     'Builds bottom-up revenue forecasts across the full patient journey: Total Population > Incidence/Prevalence > '
     'Diagnosis Rate > Treatment Rate > Eligibility Criteria > Class Share > Product Share > Annual Cost > Net Sales. '
     'Fully configurable funnel depth.'),
    ('Multi-Country & Multi-Indication Support',
     'Covers United States, Germany, United Kingdom, France, Japan, China, Canada, Italy, and Spain with '
     'pre-populated population and rebate benchmarks. Indication templates include Oncology, Rheumatoid Arthritis, '
     'Multiple Sclerosis, Type 2 Diabetes, Alzheimer Disease, and Heart Failure.'),
    ('S-Curve Adoption Modeling',
     'Market share ramps modeled via configurable S-curves: set starting share, time to peak, and peak year. '
     'Automatically calculates year-by-year patient and revenue trajectories.'),
    ('Scenario Comparison',
     'Run and compare multiple named scenarios side-by-side (base, bull, bear). '
     'Instantly see how assumption changes impact peak year revenue.'),
    ('Evidence-Backed Assumptions',
     'Every default assumption ships with a literature rationale and source URL. Teams can audit, override, and '
     'document their own rationale for regulatory or investor submissions.'),
    ('Interactive Visualizations',
     'Revenue waterfall, patient funnel chart, multi-year revenue trajectory, and market share ramp-up, '
     'all built with interactive Plotly charts exportable as images.'),
    ('Export & Integration',
     'Download forecast results as CSV or Excel. Structured JSON output enables integration '
     'with downstream BI tools and brand planning systems.'),
]

for feat_name, feat_desc in features:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run_n = p.add_run(feat_name + '   ')
    run_n.bold = True; run_n.font.size = Pt(11); run_n.font.color.rgb = NAVY
    run_d = p.add_run(feat_desc)
    run_d.font.size = Pt(10.5); run_d.font.color.rgb = DARK

doc.add_page_break()

# ═══════════════════════════════════════════════
# 4. DEMO WALKTHROUGH
# ═══════════════════════════════════════════════
add_heading('4. Recommended Demo Flow', size=17)
add_divider()
add_body(
    'Use the following step-by-step script during live demos or POC sessions. '
    'Tailor the indication and product to the prospect\'s therapeutic area whenever possible.'
)

steps = [
    ('Step 1 - Product Setup (2 min)',
     'Navigate to the Product Info step. Enter: Product = Keytruda, Country = United States, '
     'Indication = NSCLC, Launch Year = 2026, Peak Year = 2031. '
     'Highlight the dropdown for multi-country selection.',
     '[Screenshot placeholder: Product Info screen]'),
    ('Step 2 - Funnel Flow Definition (2 min)',
     'Select the "AI Recommendation" preset or "Oncology" template. Show how the parameter funnel adapts '
     'automatically: Population > Incidence > Diagnosis Rate > Eligibility Criteria > Class Share > Product Share.',
     '[Screenshot placeholder: Flow Definition screen with parameter cards]'),
    ('Step 3 - Assumption Auto-Population (3 min)',
     'Click "AI Populate" to have the AI fill all parameters with benchmarked values. Walk through the rationale '
     'panel: each assumption shows its value, confidence range, and source citation. '
     'Edit one assumption manually to show override flexibility.',
     '[Screenshot placeholder: Parameter inputs with rationale panel open]'),
    ('Step 4 - AI Copilot Chat (3 min)',
     'Open the right-hand chat panel. Ask: "What is the typical diagnosis rate for NSCLC in the US?" '
     'then "How sensitive is my forecast to the product share assumption?" '
     'Show that the AI interprets results in context of the current model.',
     '[Screenshot placeholder: Chat panel with AI response]'),
    ('Step 5 - Results & Visualization (3 min)',
     'Navigate to the Results step. Walk through: (1) patient funnel waterfall showing how 335M people funnel '
     'down to ~2,761 treated patients by 2031, (2) net sales trajectory from $47M in 2026 to $323M in 2031, '
     '(3) market share S-curve ramp to 25% peak product share.',
     '[Screenshot placeholder: Revenue chart and patient funnel]'),
    ('Step 6 - Export (1 min)',
     'Download CSV / Excel. Mention JSON API output for BI integration. '
     'Offer to share the live environment for a hands-on POC.',
     '[Screenshot placeholder: Export dialog]'),
]

for title, body, screenshot in steps:
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after  = Pt(2)
    run_t = p_title.add_run(title)
    run_t.bold = True; run_t.font.size = Pt(11); run_t.font.color.rgb = GOLD
    add_body(body)
    add_screenshot_placeholder(screenshot)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 5. OBJECTION HANDLING
# ═══════════════════════════════════════════════
add_heading('5. Objection Handling Guide', size=17)
add_divider()

objections = [
    ('We already have an Excel model that works fine.',
     'Most of our customers said the same thing before we showed them a side-by-side. The question isn\'t whether '
     'your model works - it\'s how long it takes to update, how easy it is for someone else to audit, and whether '
     'every assumption has a documented source. Can I show you how Forecasting Accelerator compares on your own product?'),
    ('Our team isn\'t technical enough for an AI tool.',
     'The entire interface is conversational - you type questions in plain English and the AI explains every '
     'parameter. If your team can fill in a spreadsheet, they can use this. We\'ve trained commercial analysts '
     'with zero data science background in under 30 minutes.'),
    ('We\'re concerned about data security and IP.',
     'No proprietary patient data leaves your environment. The AI copilot uses published literature and your '
     'own inputs - it does not share your assumptions with third parties. We can walk through our security '
     'architecture and sign an NDA before any POC.'),
    ('The assumptions feel like black boxes.',
     'Every single assumption is fully transparent - value, confidence range, and a PubMed or real-world data '
     'citation. You can override any value and document your own rationale. Reviewers and payers see everything.'),
    ('We don\'t have budget right now.',
     'Understood. Let\'s start with a no-cost 2-week POC on one of your active products. If the team saves even '
     'two weeks of analyst time per forecast cycle, the ROI conversation becomes much easier. '
     'What\'s the next forecast your team is working on?'),
    ('We already use [Competitor] for forecasting.',
     'We respect that relationship. The key question is whether your team can run scenarios interactively '
     'without filing a request with a modeler, and whether every assumption is traceable to a source. '
     'Most tools require expert configuration for every new product. Ours is ready out of the box. '
     'Worth a 30-minute comparison?'),
]

for obj, resp in objections:
    p_obj = doc.add_paragraph()
    p_obj.paragraph_format.space_before = Pt(10)
    p_obj.paragraph_format.space_after  = Pt(2)
    run_obj = p_obj.add_run('Objection: "' + obj + '"')
    run_obj.bold = True; run_obj.italic = True
    run_obj.font.size = Pt(10.5); run_obj.font.color.rgb = RED

    p_resp = doc.add_paragraph()
    p_resp.paragraph_format.space_after = Pt(4)
    shade_paragraph(p_resp, 'F0F7F0')
    run_resp = p_resp.add_run('  Response: ' + resp)
    run_resp.font.size = Pt(10.5); run_resp.font.color.rgb = GREEN_LIGHT

doc.add_page_break()

# ═══════════════════════════════════════════════
# 6. COMPETITIVE POSITIONING
# ═══════════════════════════════════════════════
add_heading('6. Competitive Positioning', size=17)
add_divider()
add_body('Use the following table when a prospect mentions competitors or compares tools.')

comp_rows = [
    ('Capability', 'Chryselys FA', 'Traditional Excel', 'Generic BI Tools', 'Legacy Pharma Models'),
    ('Time to first forecast', 'Under 30 minutes', '2-6 weeks', 'Requires data team', '4-8 weeks'),
    ('AI-assisted assumptions', 'Built-in, literature-cited', 'None', 'None', 'None'),
    ('Conversational AI copilot', 'Yes - plain English Q&A', 'No', 'No', 'No'),
    ('Evidence-backed rationale', 'Every assumption sourced', 'Manual research', 'No', 'Varies'),
    ('Multi-country support', '9 countries out of box', 'Manual setup', 'Depends', 'Limited'),
    ('Indication templates', '6 pre-built templates', 'Build from scratch', 'None', 'Limited'),
    ('Scenario comparison', 'Instant, side-by-side', 'Copy/paste sheets', 'Limited', 'Limited'),
    ('Non-technical user access', 'Full - no coding needed', 'Partial', 'Low', 'Low'),
    ('Audit trail / sources', 'Full with URLs', 'None', 'None', 'Varies'),
]

comp_tbl = doc.add_table(rows=len(comp_rows), cols=5)
comp_tbl.style = 'Table Grid'
comp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
c_widths = [Inches(1.7), Inches(1.5), Inches(1.1), Inches(1.1), Inches(1.1)]
for i, row_data in enumerate(comp_rows):
    row = comp_tbl.rows[i]
    for j, (cell, w) in enumerate(zip(row.cells, c_widths)):
        cell.width = w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(row_data[j])
        run.font.size = Pt(8.5)
        if i == 0:
            run.bold = True; run.font.color.rgb = WHITE
            shade_cell(cell, '1A4F72')
        elif j == 1:
            run.bold = True; run.font.color.rgb = GREEN_DARK
            shade_cell(cell, 'E8F5EE')
        else:
            run.font.color.rgb = DARK
            shade_cell(cell, 'F5F6F8' if i % 2 == 0 else 'FFFFFF')
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 7. SAMPLE FORECAST OUTPUT
# ═══════════════════════════════════════════════
add_heading('7. Sample Forecast Output - Keytruda / NSCLC (US)', size=17)
add_divider()
add_body(
    'Use this real forecast output during demos to show the credibility and detail of '
    'Forecasting Accelerator results. All numbers are generated by the AI from published literature assumptions.'
)

add_subheading('Key Assumptions')
for b in [
    'US Total Population: 335,000,000 (2026 baseline, 0.5% annual growth)',
    'NSCLC Incidence Rate: 1.005% per year',
    'Diagnosis Rate: 20% (National Lung Screening Trial)',
    'PD-1 Eligibility Criteria: 40%',
    'Annual Cost per Patient: $150,000',
    'Discount / Rebate Rate: 22%',
    'Time to Peak Share: 5 years (S-curve adoption, launch 2026, peak 2031)',
]:
    add_bullet(b)

add_subheading('Revenue Trajectory')
rev_rows = [
    ('Year', 'Treated Patients', 'Product Share', 'Gross Sales', 'Net Sales', 'YoY Revenue Growth'),
    ('2026', '404', '3.0%', '$60.6M', '$47.3M', 'Launch year'),
    ('2027', '760', '5.8%', '$114.1M', '$89.0M', '+88%'),
    ('2028', '1,368', '10.8%', '$205.1M', '$160.0M', '+80%'),
    ('2029', '2,045', '17.2%', '$306.8M', '$239.3M', '+50%'),
    ('2030', '2,521', '22.2%', '$378.1M', '$294.9M', '+23%'),
    ('2031 (Peak)', '2,761', '25.0%', '$414.2M', '$323.1M', '+10%'),
]
rev_tbl = doc.add_table(rows=len(rev_rows), cols=6)
rev_tbl.style = 'Table Grid'
rev_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
r_widths = [Inches(0.9), Inches(1.1), Inches(0.9), Inches(1.2), Inches(1.0), Inches(1.1)]
for i, row_data in enumerate(rev_rows):
    row = rev_tbl.rows[i]
    for j, (cell, w) in enumerate(zip(row.cells, r_widths)):
        cell.width = w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(row_data[j])
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True; run.font.color.rgb = WHITE
            shade_cell(cell, '1A4F72')
        elif i == len(rev_rows) - 1:
            run.bold = True; run.font.color.rgb = DARK
            shade_cell(cell, 'FFF3CD')
        else:
            run.font.color.rgb = DARK
            shade_cell(cell, 'F5F6F8' if i % 2 == 0 else 'FFFFFF')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

add_body('')
p_note = doc.add_paragraph()
run_note = p_note.add_run(
    'Note: Peak net sales of $323M in 2031 modeled at 25% product share within ~4% PD-1 class share '
    'of the eligible NSCLC market in the United States.'
)
run_note.font.size = Pt(9); run_note.italic = True; run_note.font.color.rgb = GREY

add_screenshot_placeholder('[Screenshot placeholder: Revenue trajectory chart - line chart 2026-2031 showing net sales ramp from $47M to $323M]')
add_screenshot_placeholder('[Screenshot placeholder: Patient funnel waterfall - 335M population funneling to 2,761 treated patients]')

doc.add_page_break()

# ═══════════════════════════════════════════════
# 8. DISCOVERY QUESTIONS
# ═══════════════════════════════════════════════
add_heading('8. Discovery Questions', size=17)
add_divider()
add_body('Use these questions in initial calls to qualify the opportunity and tailor your pitch.')

disc = {
    'Workflow & Process': [
        'How do you currently build your commercial forecasts: Excel, a licensed tool, or external consultants?',
        'How long does a typical forecast cycle take from kickoff to final output?',
        'Who owns the forecast model: one person, or is it a shared resource?',
        'How many products or markets are you forecasting right now?',
    ],
    'Pain Points': [
        'What slows you down the most in your current forecasting process?',
        'When a senior leader asks "where did this number come from?", how easy is it to answer?',
        'How do you currently handle scenario analysis: separate model files or in-tool?',
        'How much analyst time goes into literature research versus actual modeling?',
    ],
    'Stakeholders & Decision': [
        'Who reviews and signs off on your forecast methodology?',
        'Is this decision owned by commercial operations, finance, or another function?',
        'Are you evaluating any other tools at the moment?',
        'What does your ideal solution look like, and what would a successful POC prove?',
    ],
}

for section_title, qlist in disc.items():
    add_subheading(section_title, size=11, space_before=8)
    for q in qlist:
        add_bullet(q)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 9. CALL TO ACTION
# ═══════════════════════════════════════════════
add_heading('9. Call to Action & Closing Moves', size=17)
add_divider()

ctas = [
    ('Live Demo',
     '30-minute live session using the prospect\'s own product and indication. '
     'Offer to pre-configure the funnel before the call for maximum impact.'),
    ('2-Week POC',
     'No-cost proof of concept on one active product. Define success criteria upfront: '
     'time saved vs. current workflow, assumption quality, AI chat usefulness.'),
    ('Side-by-Side Comparison',
     'Run Forecasting Accelerator in parallel with their current model on the same product. '
     'Compare outputs, time invested, and assumption documentation.'),
    ('Stakeholder Briefing',
     'Offer a 45-minute briefing for the VP/C-Suite sponsor covering ROI model, '
     'security architecture, and implementation timeline.'),
    ('Reference Call',
     'Connect the prospect with a reference customer in a similar therapeutic area '
     'or company size for peer validation.'),
]

for cta_title, cta_desc in ctas:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run_t = p.add_run(cta_title + '   ')
    run_t.bold = True; run_t.font.size = Pt(11); run_t.font.color.rgb = NAVY
    run_d = p.add_run(cta_desc)
    run_d.font.size = Pt(10.5); run_d.font.color.rgb = DARK

add_divider()

p_close = doc.add_paragraph()
p_close.paragraph_format.space_before = Pt(20)
p_close.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_cl = p_close.add_run('Chryselys Forecasting Accelerator  |  Built for Pharma Commercial Teams')
run_cl.font.size = Pt(10); run_cl.bold = True; run_cl.font.color.rgb = NAVY

p_close2 = doc.add_paragraph()
p_close2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_cl2 = p_close2.add_run('Internal Use Only - Sales & Business Development')
run_cl2.font.size = Pt(9); run_cl2.italic = True; run_cl2.font.color.rgb = GREY

out = r'c:\Users\SatyamKumar\Downloads\Forecasting Agent\Chryselys_Forecasting_Accelerator_Sales_Playbook.docx'
doc.save(out)
print('Saved:', out)
